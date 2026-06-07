"""
build_report_docx.py
=====================
Generates a formatted Word document (PROJECT_REPORT.docx) of the project report,
including an embedded dashboard image. Run:  python build_report_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
NAVY = RGBColor(0x1F, 0x3A, 0x5F)

doc = Document()

# ---- Page setup: US Letter, 1-inch margins ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for side in ("left", "right", "top", "bottom"):
    setattr(sec, f"{side}_margin", Inches(1))

# ---- Default font ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


# ---- Helpers --------------------------------------------------------------
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def para(text="", bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold, run.italic = bold, italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def mono(lines):
    for ln in lines:
        p = doc.add_paragraph()
        run = p.add_run(ln)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):                      # bold header row
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        _shade(cell, "DCE6F1")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_footer_page_numbers():
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Sales & Revenue Intelligence Dashboard  |  Page ")
    run = p.add_run()
    for typ, txt in (("begin", None), ("instr", "PAGE"), ("end", None)):
        el = OxmlElement("w:fldChar") if typ != "instr" else OxmlElement("w:instrText")
        if typ == "instr":
            el.set(qn("xml:space"), "preserve")
            el.text = txt
        else:
            el.set(qn("w:fldCharType"), typ)
        run._r.append(el)


# ---- Title block ----------------------------------------------------------
para("Sales & Revenue Intelligence Dashboard", bold=True, size=22,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
para("Project Report", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Dataset: Sample - Superstore   |   Tools: Python, pandas, statsmodels, matplotlib",
     italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ---- 1. Abstract ----
heading("1. Abstract")
para("Retail businesses generate large volumes of transactional sales data, yet "
     "conventional sales dashboards are purely descriptive - they report historical "
     "revenue but do not explain profitability, forecast demand, detect anomalies, or "
     "recommend actions, and they fail on messy real-world data. This project presents "
     "an automated, end-to-end Sales & Revenue Intelligence pipeline built in Python "
     "that cleans raw sales data, identifies profit-draining products, forecasts future "
     "demand using a back-tested model, detects anomalous sales days, and produces "
     "actionable business recommendations along with an executive dashboard. It was "
     "validated on the real Sample - Superstore dataset (~9,994 transactions, 2015-2018) "
     "and uncovered loss-making bestsellers (Tables and Bookcases) that revenue-only "
     "dashboards hide.")

# ---- 2. Introduction ----
heading("2. Introduction")
para("Sales data is one of the most valuable assets a company owns, but raw data alone "
     "provides no insight. Most dashboards stop at showing past totals and top-selling "
     "products by revenue - which is misleading, because a product can sell heavily yet "
     "lose money due to deep discounting. Businesses also need to look forward "
     "(forecasting) and be alerted to unusual events. This project rebuilds the sales "
     "dashboard as an intelligent analytics system that goes from raw data all the way "
     "to business recommendations.")

# ---- 3. Problem Statement ----
heading("3. Problem Statement")
para("Conventional sales dashboards are purely descriptive - they report historical "
     "revenue but do not analyze profitability, forecast demand, detect anomalies, or "
     "recommend actions, and they assume clean data. There is a need for an automated "
     "analytics system that cleans raw sales data, identifies profit-draining products, "
     "forecasts demand reliably, flags anomalies, and produces actionable business "
     "recommendations.", italic=True)

# ---- 4. Objectives ----
heading("4. Objectives")
numbered([
    "Clean and validate messy real-world sales data automatically (with an audit log).",
    "Analyze profit and margin (not just revenue) to expose loss-making products.",
    "Forecast the next 6 months of sales using a back-tested, accuracy-validated model.",
    "Detect anomalous sales days automatically.",
    "Generate clear, quantified business recommendations.",
    "Produce an executive dashboard and a written report.",
])

# ---- 5. Existing System ----
heading("5. Existing System")
para("How it works: Traditional sales reporting is done manually in Excel or basic BI "
     "tools - data is exported, pivot tables and bar charts are built by hand, products "
     "are ranked by revenue, and a human eyeballs static charts of past totals.")
para("Limitations:", bold=True)
bullets([
    "Descriptive only (\"what happened\") - no forecasting or recommendations.",
    "Focuses on revenue (a vanity metric) - hides products that lose money.",
    "No automatic anomaly detection.",
    "Assumes clean data - breaks on nulls, duplicates, and errors.",
    "Manual and time-consuming; not reproducible.",
])

# ---- 6. Proposed System ----
heading("6. Proposed System")
para("An automated, end-to-end Python pipeline with six modules:")
table(["Module", "Function"], [
    ["Data Cleaning", "Fixes blank rows, duplicates, bad dates, outliers - logs every fix"],
    ["Profitability Analysis", "Computes profit & margin; finds loss-making bestsellers"],
    ["Forecasting", "Predicts next 6 months using a back-tested best model"],
    ["Anomaly Detection", "Automatically flags unusual sales days"],
    ["Recommendation Engine", "Converts findings into business actions"],
    ["Reporting", "Builds an executive dashboard and a written report"],
])

# ---- 7. System Architecture ----
heading("7. System Architecture / Workflow")
mono([
    "Raw Sales CSV (real Superstore data)",
    "      |",
    "  [1] CLEAN     -> fix blanks, duplicates, bad dates, outliers",
    "  [2] PROFIT    -> margin analysis + find loss-making products",
    "  [3] FORECAST  -> back-test 3 models, pick best, predict 6 months",
    "  [4] ANOMALY   -> flag unusual sales days (3-sigma rule)",
    "  [5] RECOMMEND -> generate business actions",
    "  [6] REPORT    -> dashboard.png + insights_report.md",
    "      |",
    "Actionable business decisions",
])

# ---- 8. Methodology ----
heading("8. Methodology / Modules")
numbered([
    "Data Cleaning (FIX #3): standardize columns, parse dates (drop invalid/blank), "
    "normalize text, remove duplicates, drop impossible quantities, cap price outliers.",
    "Profitability (FIX #2): group by category/sub-category, compute profit & margin, "
    "identify high-revenue but negative-profit products, analyze discount impact.",
    "Forecasting (FIX #1a): aggregate to monthly sales, back-test three models, pick "
    "the most accurate, retrain on all data, forecast 6 months ahead.",
    "Anomaly Detection (FIX #1b): compute daily sales, apply a rolling mean +/- 3-sigma "
    "band, flag days outside it.",
    "Recommendations (FIX #1c): rule-based logic turns the numbers into actions.",
    "Reporting: generate a 4-panel dashboard and a markdown insights report.",
])

# ---- 9. Algorithms ----
heading("9. Algorithms Used (and Why)")
table(["Technique", "Used For", "Why"], [
    ["Date coercion (errors=coerce)", "Cleaning", "Detect & drop invalid/blank dates"],
    ["IQR (Interquartile Range) rule", "Outlier removal", "Robust to extreme fat-finger values"],
    ["Group-by aggregation", "Profitability", "Standard way to summarize grouped data"],
    ["Median threshold", "Loss-maker detection", "Cleanly splits high- vs low-revenue products"],
    ["Holt-Winters Exp. Smoothing", "Forecasting", "Models level + trend + seasonality (Q4 peak)"],
    ["SARIMAX (Seasonal ARIMA)", "Forecasting", "Strong statistical model for seasonal series"],
    ["Log transform", "Forecasting", "Handles multiplicative (percentage) seasonality"],
    ["Walk-forward back-testing", "Model selection", "Measures honest out-of-sample accuracy"],
    ["MAPE (Mean Absolute % Error)", "Model scoring", "Interpretable metric to compare models"],
    ["Rolling mean +/- 3-sigma", "Anomaly detection", "Simple, explainable outlier test"],
    ["Rule-based logic", "Recommendations", "Converts numbers into business actions"],
])

# ---- 10. Dataset ----
heading("10. Dataset Description")
bullets([
    "Name: Sample - Superstore (a widely used public retail dataset, originally a Tableau sample).",
    "Source: Public GitHub mirror (leonism/sample-superstore), auto-downloaded by the code.",
    "Size: 10,800 raw rows -> 9,994 clean transactions after cleaning.",
    "Period: 2015-2018 (US retail).",
    "Key columns: Order Date, Region, Category, Sub-Category, Sales, Quantity, Discount, Profit.",
    "Categories: Furniture, Office Supplies, Technology (17 sub-categories).",
])

# ---- 11. Technologies ----
heading("11. Technologies Used")
bullets([
    "Language: Python 3",
    "Libraries: pandas & numpy (data wrangling), matplotlib (visualization), "
    "statsmodels (Holt-Winters & SARIMAX forecasting)",
    "Tools: Jupyter Notebook, Git/GitHub",
])

# ---- 12. Implementation ----
heading("12. Implementation")
bullets([
    "Single-file version: sales_dashboard.py - the whole pipeline in one runnable file.",
    "Modular version: main.py + src/ (pipeline.py, forecasting.py, superstore_loader.py, "
    "data_generator.py) - production-style structure.",
    "Notebook: notebooks/01_sales_analysis.ipynb - a narrative walkthrough.",
])

# ---- 13. Results & Output ----
heading("13. Results & Output")
para("Key Performance Indicators: Total Sales $2,297,201, Total Profit $286,397, "
     "Overall Margin 12.5%.", bold=True)
para("Headline finding - Loss-making bestsellers:")
table(["Sub-Category", "Sales", "Profit", "Margin"], [
    ["Tables", "$206,966", "-$17,725", "-8.6%"],
    ["Bookcases", "$114,880", "-$3,473", "-3.0%"],
])
para("Profit by Category:")
table(["Category", "Sales", "Profit", "Margin"], [
    ["Technology", "$836,154", "$145,455", "17.4%"],
    ["Furniture", "$742,000", "$18,451", "2.5%"],
    ["Office Supplies", "$719,047", "$122,491", "17.0%"],
])
para("Forecast - back-tested model selection (next 6 months = $361,842):")
table(["Model", "Out-of-sample MAPE"], [
    ["Holt-Winters (chosen)", "15.6%"],
    ["SARIMAX(1,1,1)(1,1,1,12)", "17.4%"],
    ["Holt-Winters (log)", "21.3%"],
])
para("Other outputs: 23 anomaly days flagged; cleaning retained 9,994 of 10,800 rows "
     "(92.5%); an executive dashboard image and a written insights report.")

# Embed the dashboard image if it exists.
img = ROOT / "outputs" / "dashboard.png"
if img.exists():
    para("Executive Dashboard:", bold=True)
    doc.add_picture(str(img), width=Inches(6.3))
    cap = doc.paragraphs[-1]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---- 14. Existing vs Proposed ----
heading("14. Existing vs Proposed System & Efficiency")
table(["Feature", "Existing System", "Proposed System"], [
    ["Data cleaning", "Manual / assumed clean", "Automated + audited (92.5% recovered)"],
    ["Metric focus", "Revenue (vanity)", "Profit & margin (finds loss-makers)"],
    ["Future view", "None", "Back-tested forecast (15.6% error)"],
    ["Monitoring", "Manual eyeballing", "Automatic anomaly detection"],
    ["Output", "Static charts", "Charts + written recommendations"],
    ["Effort", "Hours, manual", "One command, seconds"],
])
para("Efficiency: the system runs the full analysis in seconds, is fully reproducible, "
     "achieves a 15.6% out-of-sample forecast error, and automatically surfaced ~$21K of "
     "annual losses that a revenue-only dashboard would hide.")

# ---- 15. Real-World Applications ----
heading("15. Real-World Applications")
bullets([
    "Retail & e-commerce: monitor sales and profit; fix discount strategy.",
    "Supply chain & inventory: use the forecast to plan stock levels.",
    "Business & marketing analytics: decide which products/categories to push.",
    "Finance: margin and profitability analysis.",
    "Operations: anomaly detection for stock-outs, pricing bugs, or fraud.",
])

# ---- 16. Limitations & Future Work ----
heading("16. Limitations & Future Work")
bullets([
    "The forecast is univariate (sales over time only). Future work: a SARIMAX model with "
    "exogenous drivers (promotions, holidays) to push below 15.6% MAPE.",
    "Anomaly detection is statistical, not causal - it flags that a day is unusual, not why.",
    "A real-time version would use a trailing rolling window for live alerting.",
])

# ---- 17. Conclusion ----
heading("17. Conclusion")
para("This project demonstrates a complete data-analyst workflow: it cleans messy real "
     "data, questions vanity metrics to find hidden losses, forecasts demand with proper "
     "back-tested rigor, monitors for anomalies, and recommends concrete actions. It "
     "transforms a passive, backward-looking dashboard into an active decision-support "
     "tool - showing the difference between simply plotting data and genuinely analyzing it.")

# ---- 18. How to Run ----
heading("18. How to Run")
mono([
    "pip install pandas numpy matplotlib statsmodels",
    "python sales_dashboard.py     # single-file version",
    "python main.py                # modular version (real Superstore data)",
])
para("Outputs are written to the outputs/ folder (charts + insights report).")

# ---- 19. Role / Contribution ----
heading("19. Role / Contribution")
para("As the Data Analyst / Developer, responsibilities included: sourcing and loading "
     "the dataset, building the automated data-cleaning pipeline, performing profitability "
     "and margin analysis, building and back-testing the forecasting models, implementing "
     "anomaly detection, generating the dashboard and report, and writing the business "
     "recommendations and documentation.")

add_footer_page_numbers()

out = ROOT / "PROJECT_REPORT.docx"
doc.save(str(out))
print(f"Saved -> {out}")
